#!/usr/bin/env python3
"""Generate an AO-level System Security Plan from a program data file.

Merges a CMMC program data file (templates/program-data.schema.json shape,
YAML or JSON) with the assessment-objective dataset
(references/data/assessment-objectives.json) into a complete SSP with one
conformity block per assessment objective, mirroring
templates/ssp-structure.md.

Usage (from repo root):
    python3 scripts/generate_ssp.py path/to/program-data.yaml -o ssp.md
    python3 scripts/generate_ssp.py path/to/program-data.yaml --docx ssp.docx

Markdown is always supported. DOCX output requires python-docx
(pip install python-docx) and mirrors the same structure with headings
and tables.
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import date
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
AO_DATASET = REPO_ROOT / "references" / "data" / "assessment-objectives.json"

FAMILY_NAMES = {
    "AC": "Access Control", "AT": "Awareness & Training",
    "AU": "Audit & Accountability", "CM": "Configuration Management",
    "IA": "Identification & Authentication", "IR": "Incident Response",
    "MA": "Maintenance", "MP": "Media Protection",
    "PS": "Personnel Security", "PE": "Physical Protection",
    "RA": "Risk Assessment", "CA": "Security Assessment",
    "SC": "System and Communications Protection",
    "SI": "System and Information Integrity",
}
FAMILY_ORDER = ["AC", "AT", "AU", "CM", "IA", "IR", "MA", "MP", "PS", "PE", "RA", "CA", "SC", "SI"]

CONFORMITY_LABELS = {
    "met": "MET", "not-met": "NOT MET", "partially-met": "PARTIALLY MET",
    "not-applicable": "NOT APPLICABLE", "inherited": "MET (INHERITED)",
    "shared": "MET (SHARED)", "not-assessed": "NOT ASSESSED",
    None: "NOT ASSESSED",
}

REFERENCES = [
    "Federal Information Security Modernization Act (FISMA) of 2014",
    "32 CFR Part 2002, Controlled Unclassified Information",
    "48 CFR 52.204-21, Basic Safeguarding of Covered Contractor Information Systems",
    "DFARS 252.204-7012, Safeguarding Covered Defense Information and Cyber Incident Reporting",
    "NIST SP 800-171 Revision 2 and NIST SP 800-171A",
    "32 CFR Part 170, Cybersecurity Maturity Model Certification (CMMC) Program",
]

ASSET_SECTIONS = [
    ("cui", "CUI Assets"),
    ("contractor_risk_managed", "Contractor Risk Managed Assets"),
    ("security_protection", "Security Protection Assets"),
    ("specialized", "Specialized Assets"),
    ("out_of_scope", "Out-of-Scope Assets"),
]


def load_program(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    if path.suffix.lower() in (".yaml", ".yml"):
        try:
            import yaml
        except ImportError:
            sys.exit("pyyaml required for YAML input: pip install pyyaml")
        return yaml.safe_load(text)
    return json.loads(text)


def conformity(v: str | None) -> str:
    return CONFORMITY_LABELS.get(v, str(v).upper() if v else "NOT ASSESSED")


def build_markdown(program: dict, dataset: dict) -> str:
    org = program.get("organization", {})
    reqs_data = program.get("requirements", {}) or {}
    sources = {s["id"]: s for s in program.get("inheritance_sources", []) or []}
    L: list[str] = []
    a = L.append

    a(f"# System Security Plan (SSP)")
    a("")
    a(f"**Organization:** {org.get('name', 'TBD')}  ")
    a(f"**System:** {org.get('system_name', 'TBD')}  ")
    a(f"**Date:** {org.get('date', date.today().isoformat())}  ")
    a(f"**Revision:** {org.get('revision', 'Rev 0')}")
    a("")
    a("## Record of Acceptance/Approval")
    a("")
    owner = (org.get("roles") or {}).get("system_owner", {})
    a(f"System Owner: {owner.get('name', 'TBD')}, {owner.get('title', '')}")
    a("")
    a("Printed Name: ______________________  Date: ____________  Signature: ______________________")
    a("")
    a("## Purpose")
    a("")
    a(org.get("purpose") or (
        f"This System Security Plan documents how {org.get('name', 'the organization')} "
        f"securely operates {org.get('system_name', 'the in-scope system')} to protect "
        "Controlled Unclassified Information in accordance with NIST SP 800-171 Rev 2 "
        "and the CMMC Program (32 CFR Part 170). Conformity is recorded per NIST SP "
        "800-171A assessment objective."))
    a("")
    a("## Scope")
    a("")
    a(org.get("scope_narrative", "*To be completed: describe the CMMC Assessment Scope.*"))
    a("")
    a("## References")
    a("")
    for r in REFERENCES:
        a(f"- {r}")
    a("")
    a("## Roles & Responsibilities")
    a("")
    for slug, person in (org.get("roles") or {}).items():
        title = person.get("title", "")
        a(f"**{slug.replace('_', ' ').title()}**: {person.get('name', 'TBD')}"
          + (f", {title}" if title else ""))
        for resp in person.get("responsibilities", []) or []:
            a(f"- {resp}")
        a("")
    a("## System Environment")
    a("")
    a(org.get("environment_narrative", "*To be completed: describe the system environment.*"))
    a("")
    a("## Asset Types")
    a("")
    assets = program.get("assets", {}) or {}
    for key, label in ASSET_SECTIONS:
        a(f"### {label}")
        a("")
        rows = assets.get(key) or []
        if rows:
            a("| Asset | Description | Vendor |")
            a("|-------|-------------|--------|")
            for x in rows:
                a(f"| {x.get('name', '')} | {x.get('description', '')} | {x.get('vendor', '')} |")
        else:
            a("*None identified / to be completed.*")
        a("")
    a("## Networking Diagrams")
    a("")
    d = program.get("diagrams", {}) or {}
    gen_note = ""
    if program.get("topology"):
        gen_note = " (generate with scripts/generate_diagrams.py from the topology section)"
    a(f"- Figure 1, Network diagram: {d.get('network', '*attach*')}{gen_note}")
    a(f"- Figure 2, CUI flow diagram: {d.get('cui_flow', '*attach*')}{gen_note}")
    a("")
    hw = program.get("hardware_software") or []
    a("## Hardware and Software Information")
    a("")
    if hw:
        a("| Name | Type | Version | Location |")
        a("|------|------|---------|----------|")
        for x in hw:
            a(f"| {x.get('name', '')} | {x.get('type', '')} | {x.get('version', '')} | {x.get('location', '')} |")
    else:
        a("*To be completed from the asset inventory.*")
    a("")

    # Requirement families
    by_family: dict[str, list[dict]] = {}
    for r in dataset["requirements"]:
        by_family.setdefault(r["id"].split(".")[0], []).append(r)

    poam_rows = []
    for fam in FAMILY_ORDER:
        a(f"## {FAMILY_NAMES[fam]} [{fam}] Family")
        a("")
        for r in by_family[fam]:
            rid = r["id"]
            entry = reqs_data.get(rid, {}) or {}
            a(f"### {rid}: {r['name']}")
            a("")
            a(f"> {r['statement']}")
            a("")
            a(f"**REQUIREMENT CONFORMITY:** {conformity(entry.get('conformity'))}")
            a("")
            a("**ASSESSMENT OBJECTIVE(S):**")
            a("")
            obj_entries = entry.get("objectives", {}) or {}
            for letter, text in r["assessment_objectives"].items():
                oe = obj_entries.get(letter, {}) or {}
                a(f"**[{letter}]** {text}")
                a("")
                a(f"- AO CONFORMITY: {conformity(oe.get('conformity'))}")
                stmt = oe.get("statement") or "Assessment objective has not been assessed."
                a(f"- Assessment Objective Conformity Statement: {stmt.strip()}")
                inh = oe.get("inheritance")
                if inh:
                    src = sources.get(inh.get("source"), {})
                    label = f"{src.get('provider', inh.get('source'))} {src.get('cso', '')}".strip()
                    bits = [f"{inh.get('type', 'inherited').title()} from {label}"]
                    if inh.get("crm_ref"):
                        bits.append(f"CRM ref: {inh['crm_ref']}")
                    if src.get("crm_document"):
                        bits.append(f"CRM doc: {src['crm_document']}")
                    if src.get("boe_reference"):
                        bits.append(f"BoE: {src['boe_reference']}")
                    if inh.get("customer_note"):
                        bits.append(f"Customer share: {inh['customer_note']}")
                    a(f"- Inheritance: {'; '.join(bits)}")
                for ev in oe.get("evidence", []) or []:
                    link = f" ({ev['link']})" if ev.get("link") else ""
                    a(f"- Evidence: {ev.get('name', '')}{link}")
                a("")
            if entry.get("remediation_plan"):
                a(f"**Remediation plan:** {entry['remediation_plan'].strip()}")
                a("")
            p = entry.get("poam")
            if p:
                poam_rows.append((p.get("priority", ""), rid, p))
        a("")

    a("## Plans of Action & Milestones [POA&M]")
    a("")
    if poam_rows:
        a("| Priority | POA&M | Due Date | Actions |")
        a("|----------|-------|----------|---------|")
        for prio, rid, p in poam_rows:
            actions = "; ".join(p.get("actions", []) or [])
            a(f"| {prio} | {rid}: {p.get('description', '')} | {p.get('due', '')} | {actions} |")
        if program.get("conditional_status_date"):
            a("")
            a(f"Conditional CMMC Status Date: {program['conditional_status_date']}. "
              "All POA&M items must close within 180 days (32 CFR 170.21(b)).")
    else:
        a("*No open POA&M items.*")
    a("")
    a("## NIST CMVP Certificates")
    a("")
    certs = program.get("cmvp_certificates") or []
    if certs:
        a("| Certificate # | Vendor | Module Name | Standard | Status | Sunset Date | Associated Security Policy |")
        a("|---------------|--------|-------------|----------|--------|-------------|----------------------------|")
        for c in certs:
            a(f"| {c.get('certificate', '')} | {c.get('vendor', '')} | {c.get('module', '')} "
              f"| {c.get('standard', '')} | {c.get('status', '')} | {c.get('sunset', '')} "
              f"| {c.get('security_policy', '')} |")
        a("")
        a("Verify each certificate at the NIST CMVP registry before the assessment.")
    else:
        a("*To be completed: list CMVP certificates backing FIPS-validated cryptography claims (SC.L2-3.13.11).*")
    a("")
    return "\n".join(L)


def write_docx(markdown_lines: str, program: dict, dataset: dict, out: Path) -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        sys.exit("python-docx required for --docx output: pip install python-docx")
    from docx import Document

    doc = Document()
    for line in markdown_lines.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("> "):
            doc.add_paragraph(line[2:], style="Intense Quote")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            doc.add_paragraph(line)  # simple: tables kept as text rows
        elif line.strip():
            doc.add_paragraph(line.replace("**", ""))
    doc.save(out)


def main() -> int:
    ap = argparse.ArgumentParser(description="Generate an AO-level SSP")
    ap.add_argument("program_data", type=Path, help="program data file (YAML or JSON)")
    ap.add_argument("-o", "--out", type=Path, default=Path("ssp.md"), help="Markdown output path")
    ap.add_argument("--docx", type=Path, help="also write a DOCX to this path")
    args = ap.parse_args()

    if not AO_DATASET.exists():
        sys.exit(f"AO dataset not found: {AO_DATASET}")
    dataset = json.loads(AO_DATASET.read_text(encoding="utf-8"))
    program = load_program(args.program_data)
    if not isinstance(program, dict):
        sys.exit("program data must parse to a mapping (YAML or JSON object)")

    unknown = [k for k in (program.get("requirements") or {}) if k not in {r["id"] for r in dataset["requirements"]}]
    if unknown:
        sys.exit(f"unknown requirement ids in program data: {unknown}")

    md = build_markdown(program, dataset)
    args.out.write_text(md, encoding="utf-8")
    n_reqs = len(dataset["requirements"])
    n_aos = sum(len(r["assessment_objectives"]) for r in dataset["requirements"])
    print(f"wrote {args.out} ({n_reqs} requirements, {n_aos} assessment objectives)")
    if args.docx:
        write_docx(md, program, dataset, args.docx)
        print(f"wrote {args.docx}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
